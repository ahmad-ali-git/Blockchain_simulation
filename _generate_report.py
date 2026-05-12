"""
Generates the Final Project Report as a .docx file following the VU template.
Output: FinalSubmission/Final_Report.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "FinalSubmission/Final_Report.docx"
ASSETS = "FinalSubmission/ReportAssets"

doc = Document()

# ---------- styles ----------
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

for s in ("Heading 1", "Heading 2", "Heading 3"):
    h = doc.styles[s]
    h.font.name = "Times New Roman"
    h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
doc.styles["Heading 1"].font.size = Pt(20)
doc.styles["Heading 2"].font.size = Pt(15)
doc.styles["Heading 3"].font.size = Pt(13)

# page margins
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)


# ---------- helpers ----------
def add_para(text, bold=False, italic=False, center=False, size=None, space_after=6):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_centered_block(lines, size=14, bold=False, space_after=6):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.bold = bold


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for r_i, row in enumerate(rows, start=1):
        cells = t.rows[r_i].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = str(val)
    if col_widths:
        for r in t.rows:
            for i, w in enumerate(col_widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph()  # spacing


def add_image(path, width_inches=5.5, caption=None):
    doc.add_picture(path, width=Inches(width_inches))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(11)


# =====================================================================
# TITLE PAGE
# =====================================================================
add_centered_block([""], space_after=20)
add_centered_block(["Final Project Report"], size=22, bold=True, space_after=8)
add_centered_block([""], space_after=8)
add_centered_block(
    ["A Simulation-Based Framework for",
     "Mitigating a 51% Attack on Blockchain Networks"],
    size=18, bold=True, space_after=6,
)
add_centered_block([""], space_after=24)
add_centered_block(["Project Supervisor"], size=13, bold=False, space_after=4)
add_centered_block(["Dr. Fouzia Jumani"], size=14, bold=True, space_after=24)
add_centered_block(["Submitted By"], size=13, bold=False, space_after=4)
add_centered_block(["Group ID: F25PROJECT5043F"], size=13, bold=True, space_after=4)
add_centered_block(["Ahmad Ali"], size=14, bold=True, space_after=2)
add_centered_block(["BC220425973"], size=12, space_after=24)
add_centered_block([""], space_after=20)
add_centered_block(
    ["Software Projects & Research Section,",
     "Department of Computer Sciences,",
     "Virtual University of Pakistan"],
    size=12, space_after=4,
)
add_page_break()


# =====================================================================
# CERTIFICATE
# =====================================================================
doc.add_heading("CERTIFICATE", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    "This is to certify that Ahmad Ali (BC220425973) has worked on and "
    "completed his Software Project at Software & Research Projects Section, "
    "Department of Computer Sciences, Virtual University of Pakistan in "
    "partial fulfillment of the requirement for the degree of BS in Computer "
    "Sciences under my guidance and supervision."
)
doc.add_paragraph(
    "In my opinion, it is satisfactory and up to the mark and therefore "
    "fulfills the requirements of BS in Computer Sciences."
)
doc.add_paragraph()
add_para("Supervisor / Internal Examiner", bold=True)
add_para("Dr. Fouzia Jumani")
add_para("Supervisor,")
add_para("Software Projects & Research Section,")
add_para("Department of Computer Sciences,")
add_para("Virtual University of Pakistan")
add_para("___________________ (Signature)")
doc.add_paragraph()
add_para("External Examiner / Subject Specialist", bold=True)
add_para("___________________")
add_para("___________________ (Signature)")
doc.add_paragraph()
add_para("Accepted By: _____________")
add_para("(For office use)", italic=True, size=10)
add_page_break()


# =====================================================================
# EXORDIUM
# =====================================================================
doc.add_heading("EXORDIUM", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_centered_block([
    "In the name of Allah, the Compassionate, the Merciful.",
    "",
    "Praise be to Allah, Lord of Creation,",
    "The Compassionate, the Merciful,",
    "King of Judgment-day!",
    "",
    "You alone we worship, and to You alone we pray for help,",
    "Guide us to the straight path,",
    "The path of those who You have favored,",
    "Not of those who have incurred Your wrath,",
    "Nor of those who have gone astray.",
], size=13, space_after=2)
add_page_break()


# =====================================================================
# DEDICATION
# =====================================================================
doc.add_heading("DEDICATION", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(
    "I dedicate this work to my parents, who have stood beside me through "
    "every long night of study, and to my teachers at the Virtual University "
    "of Pakistan whose patience and feedback gave shape to my ideas. Without "
    "their support this project would not exist."
).italic = True
add_page_break()


# =====================================================================
# ACKNOWLEDGEMENT
# =====================================================================
doc.add_heading("ACKNOWLEDGEMENT", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    "All praise is due to Allah, who granted me the strength and patience "
    "to finish this project."
)
doc.add_paragraph(
    "I would like to express my sincere thanks to my supervisor, Dr. Fouzia "
    "Jumani, for her continuous guidance throughout the project. Her "
    "feedback, especially the suggestion to derive each miner's hashrate "
    "from the number of nodes it runs and to provide a graphical interface "
    "for the simulation, shaped the direction of the work in important ways."
)
doc.add_paragraph(
    "I also want to thank the academic staff at the Virtual University of "
    "Pakistan for designing a final year project that pushed me to combine "
    "what I had learned in software engineering, data structures, and "
    "cryptography into a single working system."
)
doc.add_paragraph(
    "Finally, my thanks go to the open source community behind Python, "
    "Streamlit, Matplotlib and the wider blockchain research community "
    "whose papers and reference implementations made this project possible."
)
add_page_break()


# =====================================================================
# PREFACE
# =====================================================================
doc.add_heading("PREFACE", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    "When I first read about the 51% attack in a blockchain article, what "
    "surprised me most was how simple the attack idea is and how serious "
    "the consequences are. A miner who can produce more than half of the "
    "blocks can quietly build a private chain and then publish it later, "
    "rewriting recent history. For small Proof-of-Work networks this is "
    "not a theoretical risk; it has happened in practice."
)
doc.add_paragraph(
    "This project is my attempt to study that attack inside a controlled "
    "environment and to test a defense against it. I built a small "
    "Proof-of-Work blockchain in Python, taught a malicious miner how to "
    "fork the chain, and then implemented a two-part defense based on a "
    "consecutive-block limit and a UTXO comparison. The result is a "
    "simulation that can be run through a graphical interface or from the "
    "command line and that produces measurable evidence of how the defense "
    "performs."
)
doc.add_paragraph(
    "The report is divided into three chapters following the standard "
    "Virtual University format. Chapter 1 covers requirements analysis. "
    "Chapter 2 describes the design. Chapter 3 documents the implementation, "
    "testing and results."
)
add_page_break()


# =====================================================================
# TABLE OF CONTENTS (manual)
# =====================================================================
doc.add_heading("TABLE OF CONTENTS", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_entries = [
    ("CHAPTER 1: GATHERING AND ANALYZING INFO", ""),
    ("    1.1 Introduction", ""),
    ("    1.2 Purpose", ""),
    ("    1.3 Scope", ""),
    ("    1.4 Definitions, Acronyms and Abbreviations", ""),
    ("    1.5 Project Requirements", ""),
    ("        1.5.1 Functional Requirements", ""),
    ("        1.5.2 Non-Functional Requirements", ""),
    ("    1.6 Use Cases and Usage Scenarios", ""),
    ("        1.6.1 Use Case Diagrams", ""),
    ("        1.6.2 Usage Scenarios", ""),
    ("    1.7 Development Methodology", ""),
    ("        1.7.1 Chosen Methodology", ""),
    ("        1.7.2 Reasons for Chosen Methodology", ""),
    ("        1.7.3 Work Plan (Gantt Chart)", ""),
    ("        1.7.4 Project Schedule", ""),
    ("CHAPTER 2: DESIGNING THE PROJECT", ""),
    ("    2.1 Introduction", ""),
    ("    2.2 Purpose", ""),
    ("    2.3 Scope", ""),
    ("    2.4 Definitions, Acronyms and Abbreviations", ""),
    ("    2.5 Architectural Representation", ""),
    ("    2.6 Sequence Diagrams", ""),
    ("    2.7 Class Diagram", ""),
    ("    2.8 Database Model", ""),
    ("    2.9 Graphical User Interfaces", ""),
    ("CHAPTER 3: DEVELOPMENT", ""),
    ("    3.1 Development Plan", ""),
    ("    3.2 Implementation Details", ""),
    ("    3.3 Testing", ""),
    ("    3.4 Results and Analysis", ""),
    ("    3.5 Limitations and Future Work", ""),
    ("REFERENCES", ""),
    ("APPENDIX", ""),
]
for entry, _ in toc_entries:
    p = doc.add_paragraph(entry)
    p.paragraph_format.space_after = Pt(2)
add_page_break()


# =====================================================================
# CHAPTER 1 TITLE PAGE
# =====================================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
add_centered_block(["CHAPTER 1"], size=28, bold=True, space_after=18)
add_centered_block(["Gathering and Analyzing Info"], size=22, bold=True)
add_page_break()


# =====================================================================
# CHAPTER 1
# =====================================================================
doc.add_heading("Chapter 1: Gathering and Analyzing Info", level=1)

doc.add_heading("1.1 Introduction", level=2)
doc.add_paragraph(
    "Public blockchain networks built on Proof-of-Work, such as Bitcoin and "
    "many of its earlier variants, depend on a simple rule: the longest "
    "valid chain wins. As long as honest miners control the majority of the "
    "network's hash power, this rule produces a single agreed history of "
    "transactions. The trouble starts when one miner, or a colluding group, "
    "controls more than half of the hash power. From that moment they can "
    "out-mine everyone else over the long run, build a secret private chain, "
    "and later publish it to replace the public one. Transactions that "
    "looked confirmed get reversed. Coins that were already spent can be "
    "spent again. This is what is known in the literature as a 51% attack, "
    "and several smaller Proof-of-Work coins have already been hit by it."
)
doc.add_paragraph(
    "This project does two things. First, it builds a working simulation of "
    "the attack so that the mechanics can be studied in a safe and "
    "repeatable way. Second, it implements a defense called Safe Mode "
    "Detection that combines a consecutive-block limit with a UTXO "
    "comparison step, and it measures how well the defense performs across "
    "many different attacker strengths."
)
doc.add_paragraph(
    "The simulation is written in Python. It supports both a command-line "
    "batch mode for collecting statistics and a Streamlit-based graphical "
    "interface for live demonstration. The hashrate of every miner is "
    "computed from the number of mining nodes that miner owns, which means "
    "the network behaves more like a real mining pool environment than a "
    "static configuration where someone simply types in a number."
)

doc.add_heading("1.2 Purpose", level=2)
doc.add_paragraph(
    "The purpose of this project is to give a hands-on study of how a 51% "
    "attack works on a small Proof-of-Work blockchain and to test whether "
    "the proposed Safe Mode Detection defense is able to prevent it. The "
    "system is not meant to plug into a real cryptocurrency network. It is "
    "an academic tool aimed at students, researchers and practitioners who "
    "want to understand the attack in concrete terms and who want to "
    "experiment with defensive ideas before they go anywhere near "
    "production code."
)
doc.add_paragraph(
    "A second purpose is to provide a clean, well-documented Python "
    "codebase that other VU students can extend. The code is organised into "
    "small modules: blockchain primitives, miner behaviour, defense logic "
    "and a thin user interface layer. Each module can be modified or "
    "replaced without breaking the rest of the system."
)

doc.add_heading("1.3 Scope", level=2)
doc.add_paragraph(
    "The scope of the project is limited in some important ways and these "
    "limits are deliberate. The simulator models a single Proof-of-Work "
    "blockchain with one or more honest miners and exactly one malicious "
    "miner. Mining itself is modeled probabilistically: at every step the "
    "system selects a miner using weighted random sampling where the "
    "weights are the hashrate shares. This abstraction is common in "
    "academic blockchain simulators (for example BlockSim) because it lets "
    "the researcher focus on consensus behavior without having to model "
    "every single hash attempt."
)
doc.add_paragraph(
    "The project does not model network propagation delay, message loss "
    "between nodes, or geographic distribution. It does not implement "
    "wallet software, mining pool protocols, or transaction fee markets. "
    "Difficulty is set as a configurable constant rather than being "
    "adjusted dynamically as in real Bitcoin."
)
doc.add_paragraph(
    "The attacker's behavior is also restricted to the most studied "
    "attack pattern, which is to fork the chain in secret, mine on the "
    "private fork, and publish only when the private chain is longer than "
    "the public one. Selfish mining and eclipse attacks are not in scope "
    "for this project."
)

doc.add_heading("1.4 Definitions, Acronyms and Abbreviations", level=2)
add_table(
    ["Term", "Meaning"],
    [
        ("PoW", "Proof of Work, the consensus mechanism used by Bitcoin "
                "and several other blockchains. Miners must find a hash "
                "below a target value by iterating a nonce."),
        ("51% Attack", "An attack in which a single miner or colluding "
                       "group controls the majority of the network's hash "
                       "power and uses that power to rewrite recent "
                       "blockchain history."),
        ("Hashrate", "The fraction of total network computing power "
                     "controlled by a miner. In this project it is derived "
                     "from the number of nodes the miner runs."),
        ("Node", "A single mining unit owned by a miner. The more nodes a "
                 "miner has, the larger its share of the total hashrate."),
        ("UTXO", "Unspent Transaction Output. A record of coins that have "
                 "been received and not yet spent. The UTXO set tracks all "
                 "such records."),
        ("LPC", "Long Private Chain defense. Rejects any block that would "
                "create six or more consecutive blocks by the same miner."),
        ("Safe Mode Detection",
         "Combined defense made up of LPC and a UTXO comparison check that "
         "is applied when a competing chain is published."),
        ("Genesis Block", "The very first block of the chain, with index 0 "
                          "and no real previous hash."),
        ("Double Spending", "Spending the same coins twice by reorganizing "
                            "the chain to remove the original transaction."),
        ("Streamlit", "A Python framework for building data-driven web "
                      "applications using only Python code."),
        ("VU", "Virtual University of Pakistan."),
    ],
    col_widths=[1.4, 4.5],
)

doc.add_heading("1.5 Project Requirements", level=2)
doc.add_paragraph(
    "The requirements were collected by studying the 51% attack literature, "
    "by reading the design document of the BlockSim simulator and by "
    "discussing the project goals with the supervisor. They are split into "
    "functional requirements (what the system must do) and non-functional "
    "requirements (how well it must do it)."
)

doc.add_heading("1.5.1 Functional Requirements", level=3)
add_table(
    ["ID", "Requirement"],
    [
        ("FR1", "The system shall initialize a Proof-of-Work blockchain "
                "with a configurable number of honest miners and exactly "
                "one malicious miner. Each miner's hashrate must be "
                "computed dynamically from the number of mining nodes it "
                "owns relative to the total number of nodes on the "
                "network."),
        ("FR2", "The system shall maintain a UTXO-style balance set, "
                "support transactions between users, and allow snapshots "
                "of the UTXO set so that two competing chains can be "
                "compared."),
        ("FR3", "The system shall simulate a baseline 51% attack under "
                "the standard longest-chain rule, with no defense in "
                "place, and record whether the attack succeeded."),
        ("FR4", "The system shall implement the Long Private Chain (LPC) "
                "defense by rejecting any block that would result in six "
                "or more consecutive blocks from the same miner."),
        ("FR5", "The system shall implement Safe Mode Detection by "
                "combining the LPC rule with a UTXO comparison step that "
                "fires when a competing chain is published, and shall log "
                "the activation event."),
        ("FR6", "The system shall be able to repeat the attack scenario "
                "with Safe Mode Detection turned on, and shall record "
                "whether the attack was prevented."),
        ("FR7", "The system shall support multiple simulation runs with "
                "different parameters such as hashrate ratio, total "
                "blocks and attack start time, and shall store the "
                "outcome of each run."),
        ("FR8", "The system shall produce charts that compare attack "
                "success rates with and without the defense, the "
                "distribution of chain lengths, and the contribution of "
                "each miner."),
        ("FR9", "The system shall log every important simulation event, "
                "including the start of the attack, every Safe Mode "
                "activation, every block rejection, and the final result. "
                "The logs shall be exportable to CSV from the graphical "
                "interface."),
    ],
    col_widths=[0.6, 5.4],
)

doc.add_heading("1.5.2 Non-Functional Requirements", level=3)
add_table(
    ["ID", "Requirement"],
    [
        ("NFR1 Performance", "A single simulation of thirty blocks should "
                             "complete in under five seconds on a normal "
                             "laptop. Batch runs of one hundred "
                             "simulations should complete in under five "
                             "minutes."),
        ("NFR2 Reliability", "Given the same input parameters and the "
                             "same random seed, the simulation must "
                             "produce the same results. The system must "
                             "never produce a chain whose hash links are "
                             "broken."),
        ("NFR3 Maintainability",
         "The code must be split into small, focused modules. Functions "
         "longer than fifty lines should be avoided. Important non-obvious "
         "logic must be commented."),
        ("NFR4 Usability", "The graphical interface must be operable by "
                           "someone who has never seen the project "
                           "before. All sliders must show their current "
                           "value and meaningful labels."),
        ("NFR5 Portability", "The system must run on Windows, Linux and "
                             "macOS as long as Python 3.10 or later is "
                             "installed. Only open source libraries are "
                             "to be used."),
        ("NFR6 Security",
         "Honest and malicious behavior must be clearly separated in the "
         "code. The documentation must make clear at every point that the "
         "system is a simulation and is not connected to any real "
         "cryptocurrency network."),
    ],
    col_widths=[1.6, 4.4],
)

doc.add_heading("1.6 Use Cases and Usage Scenarios", level=2)
doc.add_paragraph(
    "The system has a single primary user, who I will call the System "
    "Administrator. This is the person running the simulation, typically a "
    "student, a researcher or the project supervisor during evaluation. "
    "The four main use cases are listed below."
)

doc.add_heading("1.6.1 Use Case Diagram", level=3)
doc.add_paragraph(
    "The use case diagram below shows the relationship between the System "
    "Administrator and the four core use cases of the system."
)

# build a simple use case diagram with matplotlib
import matplotlib.pyplot as plt
from matplotlib.patches import Ellipse, Rectangle
import os
fig, ax = plt.subplots(figsize=(8, 5))
ax.set_xlim(0, 10)
ax.set_ylim(0, 6)
ax.axis("off")
# actor (stick figure)
ax.plot(1, 3, "o", markersize=22, color="black")
ax.plot([1, 1], [2.8, 1.8], "k-", linewidth=2)
ax.plot([0.6, 1.4], [2.4, 2.4], "k-", linewidth=2)
ax.plot([1, 0.6], [1.8, 1.0], "k-", linewidth=2)
ax.plot([1, 1.4], [1.8, 1.0], "k-", linewidth=2)
ax.text(1, 0.5, "System\nAdministrator", ha="center", fontsize=10, fontweight="bold")
# system boundary
ax.add_patch(Rectangle((3.5, 0.6), 6, 4.5, linewidth=2,
                       edgecolor="black", facecolor="none"))
ax.text(6.5, 5.3, "51% Attack Simulation System", ha="center",
        fontsize=11, fontweight="bold")
# 4 use cases as ellipses
ucs = [
    (5, 4.2, "UC-01\nConfigure\nSimulation"),
    (8, 4.2, "UC-02\nRun Baseline\nSimulation"),
    (5, 1.8, "UC-03\nRun Protected\nSimulation"),
    (8, 1.8, "UC-04\nAnalyze\nResults"),
]
for (x, y, label) in ucs:
    ax.add_patch(Ellipse((x, y), 2.2, 1.2,
                         edgecolor="black", facecolor="#ecf0f1", linewidth=1.5))
    ax.text(x, y, label, ha="center", va="center", fontsize=9)
    # connect actor to use case
    ax.plot([1.5, x - 1.1], [3, y], "k-", linewidth=0.8)
plt.tight_layout()
uc_path = f"{ASSETS}/diagram_usecase.png"
plt.savefig(uc_path, dpi=130)
plt.close()

add_image(uc_path, width_inches=5.8, caption="Figure 1.1: Use case diagram")

doc.add_heading("1.6.2 Usage Scenarios", level=3)

# UC-01
add_para("Use Case UC-01: Configure Simulation", bold=True, size=12)
add_table(
    ["Field", "Description"],
    [
        ("Use Case ID", "UC-01"),
        ("Use Case Name", "Configure Simulation"),
        ("Actor", "System Administrator"),
        ("Pre-conditions", "The application is installed and the Python "
                           "environment is active."),
        ("Main Flow",
         "1. The administrator launches the GUI with the streamlit run "
         "command.\n"
         "2. The administrator chooses the number of honest miners.\n"
         "3. The administrator sets the number of nodes for each honest "
         "miner. The system automatically computes the total nodes.\n"
         "4. The administrator sets the number of nodes for the attacker. "
         "The system shows the resulting hashrate as a live readout.\n"
         "5. The administrator selects the total number of blocks to mine, "
         "the PoW difficulty and the block at which the attack should "
         "start.\n"
         "6. The administrator decides whether to run the baseline scenario, "
         "the protected scenario, or both for comparison."),
        ("Post-conditions", "A valid configuration is in place and the "
                            "Run Simulation button is ready to be pressed."),
        ("Alternative Flows",
         "If the administrator sets the attacker's hashrate above 50 "
         "percent, the system displays a warning that a 51% attack is "
         "possible at this level."),
    ],
    col_widths=[1.5, 4.5],
)

# UC-02
add_para("Use Case UC-02: Run Baseline Simulation", bold=True, size=12)
add_table(
    ["Field", "Description"],
    [
        ("Use Case ID", "UC-02"),
        ("Use Case Name", "Run Baseline Simulation"),
        ("Actor", "System Administrator"),
        ("Pre-conditions", "Configuration has been completed (UC-01)."),
        ("Main Flow",
         "1. The administrator presses Run Simulation.\n"
         "2. The system creates miners using their configured node counts "
         "and computes the hashrate of each miner.\n"
         "3. The system initializes the genesis block and the UTXO set.\n"
         "4. For every step from one to the total block count, the system "
         "selects a miner using weighted random sampling and records the "
         "produced block on the public chain.\n"
         "5. At the configured attack-start block, the malicious miner "
         "starts a private chain copy and inserts a double-spend "
         "transaction.\n"
         "6. From this point on, whenever the malicious miner is selected, "
         "it mines on the private chain instead of the public one.\n"
         "7. After all blocks have been mined, the system checks whether "
         "the private chain is longer than the public one. If so, the "
         "attacker publishes the private chain and the attack succeeds."),
        ("Post-conditions", "The baseline result is recorded with the "
                            "final chain, the attack outcome and the "
                            "complete event log."),
    ],
    col_widths=[1.5, 4.5],
)

# UC-03
add_para("Use Case UC-03: Run Protected Simulation", bold=True, size=12)
add_table(
    ["Field", "Description"],
    [
        ("Use Case ID", "UC-03"),
        ("Use Case Name", "Run Protected Simulation"),
        ("Actor", "System Administrator"),
        ("Pre-conditions", "Configuration has been completed (UC-01)."),
        ("Main Flow",
         "1. The administrator presses Run Simulation with the Protected "
         "scenario selected.\n"
         "2. The system runs through the same sequence as UC-02 but with "
         "the LPC defense enabled.\n"
         "3. Before adding any block to the public chain, the system asks "
         "the LPC defense whether this would create six consecutive blocks "
         "from the same miner. If so, the block is rejected and the event "
         "is logged.\n"
         "4. When the malicious miner publishes its private chain at the "
         "end of the run, the system first checks whether the published "
         "chain itself contains a six-block streak by the same miner. If "
         "so, the chain is rejected on LPC grounds.\n"
         "5. If the chain passes the LPC check, the system runs the UTXO "
         "comparison. The honest UTXO snapshot is compared against the "
         "attacker's UTXO snapshot. Any difference is treated as evidence "
         "of a double spend, the chain is flagged and rejected."),
        ("Post-conditions", "The protected result is recorded together with "
                            "the count of blocks rejected by LPC."),
    ],
    col_widths=[1.5, 4.5],
)

# UC-04
add_para("Use Case UC-04: Analyze Results", bold=True, size=12)
add_table(
    ["Field", "Description"],
    [
        ("Use Case ID", "UC-04"),
        ("Use Case Name", "Analyze Results"),
        ("Actor", "System Administrator"),
        ("Pre-conditions", "At least one simulation has been run."),
        ("Main Flow",
         "1. The administrator scrolls to the Results section that "
         "appears inline after the simulation finishes.\n"
         "2. The system displays four metric cards (total blocks, blocks "
         "rejected, attacker hashrate, total nodes) and outcome status "
         "cards indicating whether each scenario's attack succeeded or "
         "was blocked.\n"
         "3. The administrator views the chain growth chart, the "
         "horizontal bar chart of blocks per miner, and the color-coded "
         "chain visualization.\n"
         "4. The administrator downloads the event logs as CSV files for "
         "later analysis."),
        ("Post-conditions", "The administrator has visual and tabular "
                            "evidence of how the defense performed."),
    ],
    col_widths=[1.5, 4.5],
)

doc.add_heading("1.7 Development Methodology", level=2)
doc.add_heading("1.7.1 Chosen Methodology", level=3)
doc.add_paragraph(
    "I followed the VU Process Model, which is a hybrid of the Waterfall "
    "and Spiral models that the SE-II course at the Virtual University "
    "recommends for final year projects. The work was broken into a fixed "
    "sequence of phases: requirements gathering, design, implementation, "
    "testing, and a final risk-and-refinement phase before delivery. After "
    "every phase I revisited the previous output to check whether anything "
    "needed updating."
)

doc.add_heading("1.7.2 Reasons for Chosen Methodology", level=3)
doc.add_paragraph(
    "I chose this approach for three reasons. First, the project goals "
    "were clear from the start: build a 51% attack simulation, add a "
    "defense, and measure the result. A clear-goal project benefits from "
    "the predictability that Waterfall provides. Second, the supervisor's "
    "feedback came in stages, and the Spiral aspect made it easy to fold "
    "her suggestions, in particular the dynamic node-based hashrate and "
    "the GUI requirement, into the project without throwing away earlier "
    "work. Third, the VU Process Model is the model I am most comfortable "
    "with, and using a familiar process kept me focused on solving the "
    "problem rather than learning a new methodology at the same time."
)

doc.add_heading("1.7.3 Work Plan (Gantt Chart)", level=3)

# generate gantt chart
import matplotlib.pyplot as plt
fig, ax = plt.subplots(figsize=(8.5, 4.5))
tasks = [
    ("Requirements Analysis", 0, 2),
    ("System Design", 1, 3),
    ("Implementation: Blockchain Core", 3, 4),
    ("Implementation: Miners and Network", 5, 3),
    ("Implementation: LPC Defense", 7, 2),
    ("Implementation: Safe Mode + UTXO", 8, 2),
    ("Refactor: Node-based Hashrate", 9, 2),
    ("Streamlit GUI", 10, 3),
    ("Testing and Result Collection", 12, 2),
    ("Final Report and Presentation", 13, 2),
]
colors = plt.cm.Set2(range(len(tasks)))
y_positions = list(range(len(tasks)))
y_positions.reverse()
for i, ((task, start, dur), col) in enumerate(zip(tasks, colors)):
    ax.barh(y_positions[i], dur, left=start, color=col, edgecolor="black")
    ax.text(start + dur / 2, y_positions[i], task, va="center", ha="center", fontsize=9)
ax.set_yticks(y_positions)
ax.set_yticklabels([""] * len(tasks))
ax.set_xlabel("Week")
ax.set_xticks(range(0, 16))
ax.set_title("Project Work Plan")
ax.grid(axis="x", linestyle="--", alpha=0.4)
plt.tight_layout()
plt.savefig(f"{ASSETS}/gantt_chart.png", dpi=130)
plt.close()
add_image(f"{ASSETS}/gantt_chart.png", width_inches=6.0,
          caption="Figure 1.2: Project work plan, week by week")

doc.add_heading("1.7.4 Project Schedule (Submission Calendar)", level=3)
add_table(
    ["Deliverable", "Submission Window"],
    [
        ("SRS Document", "Early in semester"),
        ("Design Document", "Mid semester"),
        ("Mid-project demo", "After core blockchain and attack working"),
        ("Final Application + Report + Presentation", "End of semester"),
    ],
    col_widths=[3.0, 3.0],
)

add_page_break()


# =====================================================================
# CHAPTER 2 TITLE PAGE
# =====================================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
add_centered_block(["CHAPTER 2"], size=28, bold=True, space_after=18)
add_centered_block(["Designing the Project"], size=22, bold=True)
add_page_break()


# =====================================================================
# CHAPTER 2
# =====================================================================
doc.add_heading("Chapter 2: Designing the Project", level=1)

doc.add_heading("2.1 Introduction", level=2)
doc.add_paragraph(
    "This chapter covers the architecture and design of the simulation. "
    "After the requirements were settled in Chapter 1, the next task was "
    "to decide how the system should be structured into modules, how those "
    "modules talk to one another, and what data structures they need. The "
    "design follows a three-tier pattern that keeps the user interface, the "
    "business logic and the data layer clearly separated."
)

doc.add_heading("2.2 Purpose", level=2)
doc.add_paragraph(
    "The purpose of the design is to translate the functional and "
    "non-functional requirements into a concrete blueprint that can be "
    "implemented in Python. The design needs to make the system easy to "
    "test, easy to extend (for example by adding new attack patterns or new "
    "defenses) and easy to demonstrate to a supervisor or evaluator."
)

doc.add_heading("2.3 Scope", level=2)
doc.add_paragraph(
    "The design covers the entire simulation from configuration to result "
    "presentation. It includes the chain data structure, the miner classes, "
    "the network helper that derives hashrate from node counts, the two "
    "defense components, the simulation engine that orchestrates the run, "
    "and the user interface that ties everything together."
)

doc.add_heading("2.4 Definitions, Acronyms and Abbreviations", level=2)
doc.add_paragraph(
    "All abbreviations introduced in section 1.4 carry over to this chapter "
    "without redefinition. Two extra design-time terms are useful here:"
)
add_table(
    ["Term", "Meaning"],
    [
        ("Three-tier architecture",
         "An architectural pattern that separates a system into a "
         "presentation tier, a business logic tier, and a data tier."),
        ("Event callback",
         "A function that the simulation engine calls every time something "
         "interesting happens during a run, used here to push live updates "
         "to the GUI."),
    ],
    col_widths=[2.0, 4.0],
)

doc.add_heading("2.5 Architectural Representation", level=2)
doc.add_paragraph(
    "The system is built as three loosely-coupled tiers. The presentation "
    "tier is responsible for collecting input from the user and displaying "
    "results. It contains the Streamlit GUI in app.py and the command-line "
    "interface in main.py. The business logic tier contains the simulation "
    "engine, the consensus rules, the LPC defense, the Safe Mode detector "
    "and the UTXO manager. The data tier contains the actual blockchain "
    "objects, the block data structure, the transaction objects, the UTXO "
    "set and the network of miners."
)
add_image(f"{ASSETS}/diagram_architecture.png", width_inches=5.8,
          caption="Figure 2.1: Three-tier architecture of the simulation system")
doc.add_paragraph(
    "An important property of this layout is that the simulation engine "
    "exposes a generic on_event callback. The engine itself does not know "
    "anything about Streamlit or about the command line. It just calls the "
    "callback whenever a block is mined, an attack starts, the LPC defense "
    "rejects a block, or the Safe Mode detector flags a chain. The GUI uses "
    "the callback to update its progress bar and live chart. The CLI ignores "
    "it. This separation is what allows both interfaces to share the same "
    "engine without any code duplication."
)

doc.add_heading("2.6 Sequence Diagrams", level=2)
doc.add_paragraph(
    "Two sequence diagrams capture the most important interactions. The "
    "first shows what happens during a baseline run when the attacker "
    "successfully publishes a longer private chain. The second shows the "
    "extra defense steps that fire during a protected run."
)

# generate sequence-diagram-like image
import matplotlib.pyplot as plt
def draw_sequence(actors, events, filename, title):
    fig, ax = plt.subplots(figsize=(9, 5))
    n = len(actors)
    spacing = 1.0 / (n + 1)
    positions = {a: spacing * (i + 1) for i, a in enumerate(actors)}
    # actor headers
    for a, x in positions.items():
        ax.text(x, 0.97, a, ha="center", va="top", fontsize=10, fontweight="bold",
                bbox=dict(boxstyle="round,pad=0.3", facecolor="#3498db",
                          edgecolor="black"), color="white")
        ax.plot([x, x], [0.93, 0.05], "k--", linewidth=0.6, alpha=0.5)
    y = 0.88
    step = (0.88 - 0.05) / (len(events) + 1)
    for src, dst, label in events:
        x1, x2 = positions[src], positions[dst]
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="->", lw=1.3, color="black"))
        mid = (x1 + x2) / 2
        ax.text(mid, y + 0.018, label, ha="center", fontsize=8.5)
        y -= step
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.set_title(title, fontsize=11, fontweight="bold")
    plt.tight_layout()
    plt.savefig(filename, dpi=130)
    plt.close()

draw_sequence(
    actors=["Admin", "GUI", "Simulation", "Blockchain", "Miner"],
    events=[
        ("Admin", "GUI", "set node counts, click Run"),
        ("GUI", "Simulation", "create Simulation, run_baseline()"),
        ("Simulation", "Miner", "select_miner() weighted by hashrate"),
        ("Simulation", "Blockchain", "mine_block(miner_id, tx)"),
        ("Blockchain", "Simulation", "return new block"),
        ("Simulation", "Miner", "attacker.start_private_chain()"),
        ("Simulation", "Miner", "attacker.mine_on_private_chain()"),
        ("Simulation", "Blockchain", "publish private chain (longer)"),
        ("Simulation", "GUI", "emit RESULT: attack_success=True"),
    ],
    filename=f"{ASSETS}/sequence_baseline.png",
    title="Sequence Diagram 2.1: Baseline 51% attack",
)
add_image(f"{ASSETS}/sequence_baseline.png", width_inches=6.2,
          caption="Figure 2.2: Sequence diagram for the baseline attack scenario")

draw_sequence(
    actors=["Admin", "GUI", "Simulation", "LPCDefense", "UTXOSet"],
    events=[
        ("Admin", "GUI", "click Run (Protected)"),
        ("GUI", "Simulation", "run_protected()"),
        ("Simulation", "LPCDefense", "should_reject_block(chain, miner)"),
        ("LPCDefense", "Simulation", "return True or False"),
        ("Simulation", "UTXOSet", "snapshot honest UTXOs"),
        ("Simulation", "UTXOSet", "snapshot attacker UTXOs"),
        ("Simulation", "LPCDefense", "check_chain_lpc_violation()"),
        ("Simulation", "UTXOSet", "compare_with(attacker_utxo)"),
        ("Simulation", "GUI", "emit DEFENSE: chain rejected"),
    ],
    filename=f"{ASSETS}/sequence_protected.png",
    title="Sequence Diagram 2.2: Protected run with Safe Mode",
)
add_image(f"{ASSETS}/sequence_protected.png", width_inches=6.2,
          caption="Figure 2.3: Sequence diagram for the protected run with Safe Mode")

doc.add_heading("2.7 Class Diagram", level=2)
doc.add_paragraph(
    "The class structure of the system is shown below. The diagram is "
    "simplified to keep it readable; only the most important attributes "
    "and methods of each class are listed."
)
add_image(f"{ASSETS}/diagram_classes.png", width_inches=6.2,
          caption="Figure 2.4: Class overview of the simulation system")
doc.add_paragraph(
    "The Block class holds a single block's data and is responsible for "
    "computing its own SHA-256 hash. The Blockchain class is the chain of "
    "blocks; it manages mining, validation and the linking of blocks. "
    "Transaction is a small data container, and UTXOSet keeps the running "
    "balance for every address as well as a compare_with method that the "
    "Safe Mode detector uses."
)
doc.add_paragraph(
    "Miner is the parent class for both HonestMiner and MaliciousMiner. "
    "Each miner stores its number of nodes; the Network class is a small "
    "helper that walks over all miners and computes their hashrate from "
    "the total node count. MaliciousMiner adds the private chain field "
    "and methods like start_private_chain and should_publish."
)
doc.add_paragraph(
    "LPCDefense holds the consecutive-block limit and the rejection "
    "counter. Simulation is the orchestrator; it owns the blockchain, the "
    "list of miners, the LPC defense and the UTXO set, and it exposes the "
    "two top-level methods run_baseline and run_protected."
)

doc.add_heading("2.8 Database Model", level=2)
doc.add_paragraph(
    "The project does not use a traditional database server. All data is "
    "either kept in memory during a run or written out as CSV files. This "
    "decision was deliberate: a simulation has no persistence requirement "
    "across sessions, and using a file-based approach makes the project "
    "easier to install on any machine."
)
doc.add_paragraph(
    "Even though there is no database, it is useful to think about what an "
    "equivalent relational schema would look like. The conceptual schema "
    "below is what the in-memory objects would map to if they were stored."
)
add_para("Table: blocks", bold=True)
add_table(
    ["Column", "Type", "Constraint"],
    [
        ("block_id", "INTEGER", "PRIMARY KEY"),
        ("previous_hash", "VARCHAR(64)", "NOT NULL"),
        ("timestamp", "DATETIME", "NOT NULL"),
        ("miner_id", "VARCHAR(50)", "FOREIGN KEY -> miners"),
        ("nonce", "INTEGER", "NOT NULL"),
        ("block_hash", "VARCHAR(64)", "UNIQUE"),
    ],
    col_widths=[1.8, 1.6, 2.6],
)
add_para("Table: transactions", bold=True)
add_table(
    ["Column", "Type", "Constraint"],
    [
        ("tx_id", "VARCHAR(64)", "PRIMARY KEY"),
        ("block_id", "INTEGER", "FOREIGN KEY -> blocks"),
        ("sender", "VARCHAR(100)", "NOT NULL"),
        ("receiver", "VARCHAR(100)", "NOT NULL"),
        ("amount", "DECIMAL(18,8)", "NOT NULL"),
    ],
    col_widths=[1.8, 1.6, 2.6],
)
add_para("Table: simulation_runs", bold=True)
add_table(
    ["Column", "Type", "Constraint"],
    [
        ("run_id", "INTEGER", "PRIMARY KEY AUTO_INCREMENT"),
        ("start_time", "DATETIME", "NOT NULL"),
        ("defense_enabled", "BOOLEAN", "NOT NULL"),
        ("attack_success", "BOOLEAN", "NULL"),
        ("total_blocks", "INTEGER", "NOT NULL"),
    ],
    col_widths=[1.8, 1.6, 2.6],
)
doc.add_paragraph(
    "The relationships between these tables are straightforward. A "
    "blockchain contains many blocks, a block contains many transactions, "
    "a miner mines many blocks, and a simulation run produces many event "
    "log entries. In the actual code these relationships exist as Python "
    "object references rather than foreign keys."
)

doc.add_heading("2.9 Graphical User Interfaces", level=2)
doc.add_paragraph(
    "The GUI was built with Streamlit. Streamlit lets a Python developer "
    "write a fully interactive web interface using only Python code. The "
    "interface follows a single-page vertical flow: the user configures "
    "parameters at the top, presses a centered Run button, and the "
    "simulation results appear inline below without any tab switching."
)
doc.add_paragraph(
    "The first section is the configuration bento grid. It is a "
    "three-column panel of cards: one for honest miners (stepper "
    "controls per miner with a visual node-dot grid), one for the "
    "attacker (node count, live hashrate readout, and a warning badge "
    "when the attacker crosses 50 percent), and one for simulation "
    "parameters (blocks, difficulty, LPC limit, run mode). A network "
    "topology ring diagram below the grid shows the relative node "
    "distribution between honest miners and the attacker."
)
doc.add_paragraph(
    "The second section is the live simulation view that renders "
    "immediately after the Run button is pressed. It shows a progress "
    "bar, a color-coded event log (green for mined blocks, red for "
    "attacker activity, blue for LPC rejections), and a Plotly line "
    "chart that plots the honest chain length and the attacker's private "
    "chain length over time. This makes the attack visually obvious "
    "during a demonstration."
)
doc.add_paragraph(
    "The third section is the results dashboard that appears once the "
    "simulation finishes. It contains four metric cards (total blocks, "
    "blocks rejected by LPC, attacker hashrate, and total nodes), "
    "outcome status cards (SUCCEEDED in red or BLOCKED in green for "
    "each scenario), a four-line chain growth chart for side-by-side "
    "comparison, a horizontal bar chart of blocks mined per miner, "
    "a block-by-block chain visualizer with color-coded miner badges, "
    "and CSV download buttons for both event logs."
)

add_page_break()


# =====================================================================
# CHAPTER 3 TITLE PAGE
# =====================================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
add_centered_block(["CHAPTER 3"], size=28, bold=True, space_after=18)
add_centered_block(["Development"], size=22, bold=True)
add_page_break()


# =====================================================================
# CHAPTER 3
# =====================================================================
doc.add_heading("Chapter 3: Development", level=1)

doc.add_heading("3.1 Development Plan", level=2)
doc.add_paragraph(
    "Implementation followed the order suggested by the design. I started "
    "with the lowest layer, the block and chain primitives, because nothing "
    "else can be built on top of them. After that I built the miner classes "
    "and the network helper that turns node counts into hashrates. Once a "
    "single mining cycle was working end to end, I added the malicious "
    "miner with its private chain logic. With the attack running, I added "
    "the LPC defense, then the UTXO comparison step that completes the "
    "Safe Mode detector. The Streamlit GUI came last, after the CLI version "
    "had been tested with batch runs."
)
doc.add_paragraph(
    "I worked in small commits and ran the existing simulation after every "
    "change. This caught a couple of regressions early. For example, when I "
    "first changed the constructor of Miner from accepting a hashrate to "
    "accepting num_nodes, the call sites in main.py broke. Catching this "
    "the moment it happened was a lot less painful than discovering it later."
)

doc.add_heading("3.2 Implementation Details", level=2)

doc.add_heading("3.2.1 Project Layout", level=3)
add_table(
    ["File", "Responsibility"],
    [
        ("block.py", "Defines the Block class, including SHA-256 hashing."),
        ("blockchain.py", "Defines the Blockchain class, including PoW "
                          "mining, the genesis block and chain validation."),
        ("transaction.py", "Defines the Transaction and UTXOSet classes."),
        ("miner.py", "Defines Miner, HonestMiner, MaliciousMiner and the "
                     "Network helper that derives hashrate from node count."),
        ("lpc_defense.py", "Defines the LPCDefense class with the "
                           "consecutive-block check."),
        ("simulation.py", "Defines the Simulation class that orchestrates "
                          "every run and emits events."),
        ("main.py", "Command-line batch runner."),
        ("app.py", "Streamlit graphical interface."),
        ("visualize.py", "Matplotlib helpers for charts."),
        ("requirements.txt", "Python dependency list."),
        ("README.md", "Setup and usage instructions."),
    ],
    col_widths=[1.6, 4.4],
)

doc.add_heading("3.2.2 Node-based Hashrate", level=3)
doc.add_paragraph(
    "The most important design change during implementation was moving "
    "from a hard-coded hashrate to a node-based one. In the first version "
    "of the code the malicious miner was created like this:"
)
add_para("    MaliciousMiner('Attacker', hashrate=0.55)",
         italic=True, size=11)
doc.add_paragraph(
    "The supervisor pointed out that this is not how real mining pools "
    "work. In a real network, a miner has a certain amount of hardware, "
    "and its share of the network's hashrate falls out of the ratio "
    "between its hardware and everyone else's. After the refactor the "
    "constructor takes the number of mining nodes instead, and a small "
    "Network helper class computes hashrates as part of setup:"
)
add_para("    HonestMiner('Honest_1', num_nodes=2)", italic=True, size=11)
add_para("    MaliciousMiner('Attacker', num_nodes=12)", italic=True, size=11)
add_para("    network = Network([... all miners ...])  # computes "
         "hashrate = num_nodes / total_nodes",
         italic=True, size=11)
doc.add_paragraph(
    "With this change, adding nodes to a miner increases its hashrate "
    "proportionally and decreasing them does the opposite. The total "
    "hashrate of the network is always one hundred percent and is split "
    "between miners according to their share."
)

doc.add_heading("3.2.3 Event Streaming for Live UI", level=3)
doc.add_paragraph(
    "The Simulation class accepts an optional callback called on_event. "
    "Whenever something interesting happens, the engine calls the callback "
    "with a small dictionary describing the event. The CLI passes no "
    "callback at all, so the engine simply prints to standard output. The "
    "GUI passes a callback that updates a Streamlit progress bar, appends "
    "to the live event log and pushes a new point to the chain growth "
    "chart. This is what makes the live view feel responsive. It also "
    "keeps the engine itself completely free of UI concerns."
)

doc.add_heading("3.3 Testing", level=2)
doc.add_paragraph(
    "Testing was carried out at three levels: unit, integration and "
    "scenario. Unit tests checked individual functions such as the block "
    "hashing function, the UTXO update logic and the LPC consecutive-block "
    "counter. Integration tests ran short simulations of around ten blocks "
    "to confirm that miners produced valid chains and that the chain "
    "linkage stayed correct. Scenario tests are the ones reported in the "
    "next section: full-length simulations with and without the defense, "
    "repeated many times to give a statistical picture."
)
doc.add_paragraph(
    "The most important test cases are summarised in the table below."
)
add_table(
    ["TC ID", "What is tested", "Expected result"],
    [
        ("TC-01", "Configure simulation parameters",
         "Configuration is accepted and the GUI updates the live "
         "hashrate readout."),
        ("TC-02", "Baseline 51% attack with no defense",
         "Attack succeeds in a measurable fraction of runs and the "
         "private chain replaces the honest chain when longer."),
        ("TC-03", "LPC block rejection during mining",
         "The sixth consecutive block from the same miner is rejected and "
         "an event is logged."),
        ("TC-04", "Safe Mode UTXO comparison",
         "When the attacker publishes a private chain that contains a "
         "double-spend, the UTXO comparison flags the chain as fraudulent "
         "and rejects it."),
        ("TC-05", "Visualization and CSV download",
         "All charts render correctly and the event log is downloadable "
         "as a CSV file."),
    ],
    col_widths=[0.7, 2.5, 3.0],
)

doc.add_heading("3.4 Results and Analysis", level=2)
doc.add_paragraph(
    "Two experiments were run to measure the effect of the defense. The "
    "first one fixed the attacker hashrate at about 54.5 percent (twelve "
    "attacker nodes against ten honest nodes) and ran thirty paired "
    "simulations: each pair consisted of one baseline run and one "
    "protected run with the same parameters. The second experiment swept "
    "the attacker hashrate from below 30 percent up to above 70 percent "
    "and recorded the success rate at each level."
)

add_para("Experiment 1: fixed attacker hashrate of 54.5 percent",
         bold=True)
add_image(f"{ASSETS}/chart_attack_success.png", width_inches=5.5,
          caption="Figure 3.1: Attack success rate, baseline vs protected, "
                  "30 runs at attacker hashrate 54.5 percent")
doc.add_paragraph(
    "In the baseline scenario the attacker succeeded in slightly under "
    "half of the runs, which matches the expectation that a 54 percent "
    "majority gives the attacker a real but not guaranteed advantage. With "
    "Safe Mode Detection enabled the attack success rate dropped to zero. "
    "The LPC rule alone caught most cases by rejecting blocks that would "
    "have produced a long streak from a single miner; the UTXO comparison "
    "covered the rest."
)

add_para("Experiment 2: attacker hashrate sweep", bold=True)
add_image(f"{ASSETS}/chart_hashrate_vs_success.png", width_inches=5.7,
          caption="Figure 3.2: Attack success rate at different attacker "
                  "hashrates, baseline vs protected")
doc.add_paragraph(
    "The sweep confirms the trend. As the attacker's hashrate increases, "
    "the baseline success rate climbs steeply, reaching one hundred "
    "percent at very high majorities. The protected line stays flat at "
    "zero across the entire range. Even when the attacker controlled more "
    "than 70 percent of the network, Safe Mode Detection still prevented "
    "every attack in the sample."
)

add_para("Chain length distribution", bold=True)
add_image(f"{ASSETS}/chart_chain_length_dist.png", width_inches=5.7,
          caption="Figure 3.3: Distribution of final chain length across "
                  "30 runs")
doc.add_paragraph(
    "The distribution of final chain length is similar in both scenarios. "
    "This is a useful sanity check; the defense reduces attack success "
    "without changing the fundamental progress rate of the chain. In other "
    "words, honest miners are not paying a noticeable performance cost "
    "for the protection."
)

add_para("Single-run chain growth", bold=True)
add_image(f"{ASSETS}/chart_chain_growth.png", width_inches=5.7,
          caption="Figure 3.4: Honest chain vs attacker private chain "
                  "during a single baseline run")
doc.add_paragraph(
    "The single-run chart shows what is going on inside one simulation. "
    "Until the attack starts at block ten, only the honest chain grows. "
    "After that the private chain starts climbing in parallel; if the "
    "attacker's hashrate is high enough, the private chain pulls ahead and "
    "wins. If not, the honest chain wins and the attack fails."
)

doc.add_heading("3.5 Limitations and Future Work", level=2)
doc.add_paragraph(
    "The simulation has limits that I want to be honest about. First, "
    "mining is modeled as weighted random selection rather than as actual "
    "competing hash work. This is a standard simplification in academic "
    "blockchain simulators but it does mean the simulation cannot model "
    "things like real hardware variance or short bursts of luck. Second, "
    "there is no model of network latency. In real Bitcoin some 51% "
    "attacks rely partly on slow propagation between nodes, and that "
    "factor cannot be reproduced here. Third, the only attack pattern "
    "implemented is the secret-fork-and-publish pattern. Real attackers "
    "have other strategies such as selfish mining and eclipse attacks "
    "that would need separate work to study."
)
doc.add_paragraph(
    "Future work could extend the project in several directions. The most "
    "obvious one is to add network latency and partition modeling, which "
    "would make the simulation more realistic and would let the defense "
    "be tested under harsher conditions. Another would be to compare the "
    "Safe Mode Detection approach with other published defenses such as "
    "checkpointing or finality gadgets. A third would be to scale the "
    "simulation up and run it on much longer chains and many more miners "
    "to see whether the defense holds at that scale."
)

add_page_break()


# =====================================================================
# REFERENCES
# =====================================================================
doc.add_heading("References", level=1)
refs = [
    "Nakamoto, S. (2008). Bitcoin: A Peer-to-Peer Electronic Cash System. "
    "Available at https://bitcoin.org/bitcoin.pdf.",
    "Eyal, I. and Sirer, E. G. (2014). Majority is not Enough: Bitcoin "
    "Mining is Vulnerable. In Financial Cryptography and Data Security.",
    "Alharby, M. and van Moorsel, A. (2019). BlockSim: A Simulation "
    "Framework for Blockchain Systems. ACM SIGMETRICS Performance "
    "Evaluation Review, 46(3), 135-138.",
    "Saad, M., Spaulding, J., Njilla, L., Kamhoua, C., Shetty, S., Nyang, "
    "D. and Mohaisen, D. (2020). Exploring the Attack Surface of "
    "Blockchain: A Systematic Overview.",
    "Bonneau, J., Miller, A., Clark, J., Narayanan, A., Kroll, J. A. and "
    "Felten, E. W. (2015). SoK: Research Perspectives and Challenges for "
    "Bitcoin and Cryptocurrencies. IEEE Symposium on Security and Privacy.",
    "Streamlit Inc. (2024). Streamlit Documentation. "
    "https://docs.streamlit.io.",
    "Python Software Foundation. (2024). Python 3 Documentation. "
    "https://docs.python.org/3/.",
    "Hunter, J. D. (2007). Matplotlib: A 2D Graphics Environment. "
    "Computing in Science and Engineering, 9(3), 90-95.",
]
for r in refs:
    p = doc.add_paragraph(r, style="List Number")
    p.paragraph_format.space_after = Pt(6)

add_page_break()


# =====================================================================
# APPENDIX
# =====================================================================
doc.add_heading("Appendix", level=1)

doc.add_heading("A. How to Run the Simulation", level=2)
doc.add_paragraph(
    "Step 1. Create a Python virtual environment and install the "
    "dependencies:"
)
add_para("    python3 -m venv venv", italic=True, size=11)
add_para("    source venv/bin/activate    (on Windows: venv\\Scripts\\activate)",
         italic=True, size=11)
add_para("    pip install -r requirements.txt", italic=True, size=11)
doc.add_paragraph("Step 2. Launch the GUI:")
add_para("    streamlit run app.py", italic=True, size=11)
doc.add_paragraph(
    "Open the URL printed in the terminal. In the configuration panel "
    "at the top of the page, set the miner node counts and simulation "
    "parameters, then press the Run Simulation button."
)
doc.add_paragraph("Step 3. Or run the CLI batch mode:")
add_para("    python main.py", italic=True, size=11)
doc.add_paragraph(
    "This runs ten paired simulations and prints a summary of attack "
    "success rates and blocks rejected by LPC."
)

doc.add_heading("B. Sample Output (CLI)", level=2)
add_para(
    "============================================================\n"
    "  FINAL ANALYSIS\n"
    "============================================================\n"
    "  Total Simulation Runs:        10\n"
    "  Baseline Attack Successes:     6/10 (60%)\n"
    "  Protected Attack Successes:    0/10 (0%)\n"
    "  Total Blocks Rejected by LPC:  14\n"
    "  Avg Blocks Rejected per Run:   1.4\n"
    "============================================================\n"
    "\n"
    "  CONCLUSION: LPC Defense significantly reduces\n"
    "  the success rate of 51% attacks!\n"
    "============================================================",
    italic=True, size=10,
)

doc.add_heading("C. Tools and Libraries Used", level=2)
add_table(
    ["Tool / Library", "Purpose"],
    [
        ("Python 3.10", "Programming language"),
        ("Streamlit", "Graphical user interface"),
        ("Matplotlib", "Plotting and chart generation"),
        ("Plotly", "Interactive charts inside the Streamlit GUI"),
        ("Pandas", "Data tabulation in the GUI and CSV export"),
        ("hashlib (stdlib)", "SHA-256 hashing for blocks"),
        ("Visual Studio Code", "Development environment"),
    ],
    col_widths=[2.0, 4.0],
)


# ---------- save ----------
doc.save(OUT)
print(f"\nReport written to {OUT}")
print(f"Pages will depend on Word's rendering, but content is complete.")
